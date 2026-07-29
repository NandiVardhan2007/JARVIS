"""
Report generation — turns structured content into a real, saved document.

Tools/document_processor.py can READ PDF/DOCX/TXT files (Q&A over them),
but nothing in this codebase could WRITE one — "generate reports" had no
home. This fills that gap with a straightforward Word/text report writer.

This is deliberately simple (title + heading/body sections, optional
bullet lists) rather than a full templating engine — VISION assembles the
content (from research, system reports, summaries, etc.) and this tool lays
it out and saves it.
"""

import json
import logging
import os
from datetime import datetime
from typing import Optional

from livekit.agents import function_tool

logger = logging.getLogger(__name__)

REPORTS_DIR = os.path.join(os.path.expanduser("~"), "Documents", "VISION", "reports")


@function_tool
async def generate_report(
    title: str,
    sections_json: str,
    output_format: str = "docx",
    output_path: Optional[str] = None,
) -> str:
    """
    Generates a formatted report document and saves it to disk.

    Args:
        title: Report title (used as the document heading and default filename).
        sections_json: A JSON array of section objects, each with:
            - 'heading' (str): Section heading text.
            - 'body' (str, optional): Paragraph text for the section.
            - 'bullets' (list of str, optional): Bullet points for the section.
            At least one of 'body' or 'bullets' should be present per section.
            Example: '[{"heading": "Summary", "body": "Q3 revenue grew 12%."},
                       {"heading": "Key Risks", "bullets": ["Supply delays", "FX exposure"]}]'
        output_format: "docx" (Word document, default) or "txt" (plain text).
        output_path: Optional absolute path to save to. Defaults to
            ~/Documents/VISION/reports/<title>.<ext>.
    """
    try:
        sections = json.loads(sections_json)
    except json.JSONDecodeError as e:
        return f"Invalid JSON in sections_json: {e}"

    if not isinstance(sections, list) or not sections:
        return "sections_json must be a non-empty JSON array of section objects."

    safe_title = "".join(c for c in title if c.isalnum() or c in " _-").strip() or "Report"
    os.makedirs(REPORTS_DIR, exist_ok=True)

    if output_format == "docx":
        path = output_path or os.path.join(REPORTS_DIR, f"{safe_title}.docx")
        return await _write_docx(title, sections, path)
    elif output_format == "txt":
        path = output_path or os.path.join(REPORTS_DIR, f"{safe_title}.txt")
        return _write_txt(title, sections, path)
    else:
        return f"Unknown output_format '{output_format}'. Use 'docx' or 'txt'."


async def _write_docx(title: str, sections: list, path: str) -> str:
    try:
        import docx
    except ImportError:
        return "python-docx is not installed. Run: pip install python-docx --break-system-packages"

    try:
        doc = docx.Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}").italic = True

        for section in sections:
            heading = section.get("heading", "").strip()
            if heading:
                doc.add_heading(heading, level=1)
            body = section.get("body", "").strip()
            if body:
                doc.add_paragraph(body)
            bullets = section.get("bullets") or []
            for bullet in bullets:
                doc.add_paragraph(str(bullet), style="List Bullet")

        doc.save(path)
        return f"Report saved to {path}."
    except Exception as e:
        logger.error(f"Failed to generate docx report: {e}")
        return f"Failed to generate report: {e}"


def _write_txt(title: str, sections: list, path: str) -> str:
    try:
        lines = [title, "=" * len(title), f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", ""]
        for section in sections:
            heading = section.get("heading", "").strip()
            if heading:
                lines.append(heading)
                lines.append("-" * len(heading))
            body = section.get("body", "").strip()
            if body:
                lines.append(body)
            bullets = section.get("bullets") or []
            for bullet in bullets:
                lines.append(f"  • {bullet}")
            lines.append("")

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return f"Report saved to {path}."
    except Exception as e:
        logger.error(f"Failed to generate text report: {e}")
        return f"Failed to generate report: {e}"
