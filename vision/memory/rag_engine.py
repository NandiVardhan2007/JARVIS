"""
Document RAG (Retrieval-Augmented Generation) & Semantic Content Search Engine for VISION.
Extracts, semantically chunks, and retrieves relevant passages from PDF, DOCX, CSV, JSON, TXT, MD,
and multi-language source code files using the Okapi BM25 ranking algorithm and multi-document directory scanning.
"""

import os
import math
import re
import csv
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from vision.tools.registry import tool
from vision.memory.working_memory import working_memory
from vision.logger import logger

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None


class DocumentRAGEngine:
    """Enterprise-grade Document text extractor, chunker, and Okapi BM25 semantic retrieval engine."""

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".json", ".xml", ".html",
        ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".cpp", ".c", ".h", ".cs",
        ".sql", ".sh", ".bat", ".ps1", ".yaml", ".yml", ".log", ".ini", ".env"
    }

    def extract_text_from_file(self, file_path: Path) -> str:
        """Extract plain text and structural content from various file formats."""
        if not file_path.exists():
            return ""

        ext = file_path.suffix.lower()

        # 1. PDF extraction
        if ext == ".pdf":
            if not pypdf:
                return "Error: pypdf not installed. Please install pypdf."
            try:
                reader = pypdf.PdfReader(str(file_path))
                text_parts = []
                for i, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t and t.strip():
                        text_parts.append(f"--- [Page {i+1}] ---\n{t.strip()}")
                return "\n\n".join(text_parts) if text_parts else "PDF is empty or scanned image only."
            except Exception as e:
                return f"Error extracting PDF text: {e}"

        # 2. DOCX extraction
        elif ext in [".docx", ".doc"]:
            if not docx:
                return "Error: python-docx not installed. Please install python-docx."
            try:
                doc = docx.Document(str(file_path))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                # Also include table content if present
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                        if row_text:
                            paragraphs.append(f"Table Row: {row_text}")
                return "\n\n".join(paragraphs)
            except Exception as e:
                return f"Error extracting DOCX text: {e}"

        # 3. CSV extraction
        elif ext == ".csv":
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    reader = csv.reader(f)
                    rows = list(reader)
                    if not rows:
                        return "CSV file is empty."
                    header = rows[0]
                    lines = [f"Columns: {', '.join(header)}"]
                    for idx, row in enumerate(rows[1:100], 1):  # Cap first 100 rows for preview
                        lines.append(f"Row {idx}: {', '.join(row)}")
                    return "\n".join(lines)
            except Exception as e:
                return f"Error reading CSV: {e}"

        # 4. JSON extraction
        elif ext == ".json":
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2)
            except Exception as e:
                return f"Error reading JSON: {e}"

        # 5. Plain text, Markdown, Logs, & Code extraction
        else:
            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
            except Exception as e:
                return f"Error reading text file: {e}"

    def chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        overlap: int = 100,
        preserve_paragraphs: bool = True
    ) -> List[Dict[str, Any]]:
        """
        Split document text into overlapping chunks with paragraph awareness and metadata.
        Returns a list of dicts: [{'text': chunk_text, 'index': i, 'word_count': n}]
        """
        if not text or not text.strip():
            return []

        if preserve_paragraphs:
            # Split by double newline or headers
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
            chunks = []
            current_words: List[str] = []
            current_chunk_idx = 1

            for p in paragraphs:
                p_words = p.split()
                if len(current_words) + len(p_words) <= chunk_size:
                    current_words.extend(p_words)
                else:
                    if current_words:
                        chunk_str = " ".join(current_words)
                        chunks.append({
                            "index": current_chunk_idx,
                            "text": chunk_str,
                            "word_count": len(current_words)
                        })
                        current_chunk_idx += 1
                        # Retain overlap from end of current chunk
                        current_words = current_words[-overlap:] if overlap < len(current_words) else []
                    
                    # If paragraph itself is larger than chunk_size, split by sliding window
                    if len(p_words) > chunk_size:
                        for i in range(0, len(p_words), chunk_size - overlap):
                            sub_chunk = " ".join(p_words[i:i + chunk_size])
                            chunks.append({
                                "index": current_chunk_idx,
                                "text": sub_chunk,
                                "word_count": len(sub_chunk.split())
                            })
                            current_chunk_idx += 1
                        current_words = []
                    else:
                        current_words.extend(p_words)

            if current_words:
                chunks.append({
                    "index": current_chunk_idx,
                    "text": " ".join(current_words),
                    "word_count": len(current_words)
                })
            return chunks
        else:
            words = text.split()
            chunks = []
            current_chunk_idx = 1
            for i in range(0, len(words), max(1, chunk_size - overlap)):
                chunk_words = words[i:i + chunk_size]
                chunks.append({
                    "index": current_chunk_idx,
                    "text": " ".join(chunk_words),
                    "word_count": len(chunk_words)
                })
                current_chunk_idx += 1
            return chunks

    def rank_passages_bm25(
        self,
        chunks: List[Dict[str, Any]],
        query: str,
        top_k: int = 3,
        k1: float = 1.5,
        b: float = 0.75
    ) -> List[Tuple[float, Dict[str, Any]]]:
        """
        Rank passages using the Okapi BM25 algorithm with exact-phrase boosting.
        """
        if not chunks or not query:
            return []

        # Tokenize query
        query_terms = [t.lower() for t in re.findall(r"\w+", query) if len(t) > 1]
        if not query_terms:
            return [(1.0, c) for c in chunks[:top_k]]

        N = len(chunks)
        # Calculate document lengths & average document length
        doc_lengths = [len(re.findall(r"\w+", c["text"])) for c in chunks]
        avgdl = sum(doc_lengths) / N if N > 0 else 1.0

        # Calculate document frequency n(q) for each query term
        doc_freq: Dict[str, int] = {}
        for term in query_terms:
            doc_freq[term] = sum(1 for c in chunks if term in c["text"].lower())

        # Compute BM25 scores
        scored_chunks: List[Tuple[float, Dict[str, Any]]] = []
        clean_query_str = " ".join(query_terms)

        for i, chunk in enumerate(chunks):
            chunk_text = chunk["text"]
            chunk_lower = chunk_text.lower()
            doc_len = doc_lengths[i]
            score = 0.0

            for term in set(query_terms):
                tf = chunk_lower.count(term)
                if tf == 0:
                    continue
                nq = doc_freq.get(term, 0)
                # BM25 IDF formulation
                idf = math.log(1.0 + (N - nq + 0.5) / (nq + 0.5))
                # BM25 TF formulation with document length normalization
                tf_norm = (tf * (k1 + 1.0)) / (tf + k1 * (1.0 - b + b * (doc_len / avgdl)))
                score += idf * tf_norm

            # Exact multi-word phrase matching bonus
            if len(query_terms) > 1 and clean_query_str in chunk_lower:
                score += 5.0

            if score > 0:
                scored_chunks.append((score, chunk))

        scored_chunks.sort(key=lambda x: x[0], reverse=True)
        return scored_chunks[:top_k]

    def search_in_document(self, file_path: Path, query: str, top_k: int = 3) -> str:
        """Retrieve the most relevant passages from a document for a query using Okapi BM25."""
        raw_text = self.extract_text_from_file(file_path)
        if not raw_text or raw_text.startswith("Error"):
            return raw_text or f"No text could be extracted from '{file_path.name}'."

        chunks = self.chunk_text(raw_text)
        if not chunks:
            return f"The document '{file_path.name}' is empty."

        ranked = self.rank_passages_bm25(chunks, query, top_k=top_k)
        if not ranked:
            top_chunks = [c["text"] for c in chunks[:top_k]]
        else:
            top_chunks = [item[1]["text"] for item in ranked]

        result = [f"📄 RAG Document Context from '{file_path.name}' ({len(chunks)} chunks analyzed):"]
        for idx, passage in enumerate(top_chunks, 1):
            result.append(f"\n[Passage {idx}]\n{passage}")

        return "\n".join(result)

    def search_in_directory(
        self,
        directory_path: Path,
        query: str,
        top_k: int = 5,
        max_files: int = 25
    ) -> str:
        """Search across multiple documents in a folder and rank relevant passages across files."""
        if not directory_path.exists() or not directory_path.is_dir():
            return f"Error: Directory '{directory_path}' does not exist or is not a directory."

        all_candidate_files = []
        for p in directory_path.rglob("*"):
            if p.is_file() and p.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                if not any(part.startswith(".") or part == "__pycache__" or part == "node_modules" for part in p.parts):
                    all_candidate_files.append(p)
                    if len(all_candidate_files) >= max_files:
                        break

        if not all_candidate_files:
            return f"No readable documents found in directory '{directory_path}'."

        all_chunks: List[Dict[str, Any]] = []
        for file_p in all_candidate_files:
            raw = self.extract_text_from_file(file_p)
            if raw and not raw.startswith("Error"):
                chunks = self.chunk_text(raw, chunk_size=400, overlap=80)
                for c in chunks:
                    c["source_file"] = file_p.name
                    c["source_path"] = str(file_p)
                    all_chunks.append(c)

        if not all_chunks:
            return f"No content could be extracted from files in '{directory_path}'."

        ranked = self.rank_passages_bm25(all_chunks, query, top_k=top_k)
        if not ranked:
            return f"No relevant content found in '{directory_path}' matching '{query}'."

        output = [f"📚 Multi-Document RAG Results from '{directory_path.name}' for query: '{query}':"]
        for rank, (score, chunk) in enumerate(ranked, 1):
            output.append(
                f"\n[Result {rank}] Source: {chunk.get('source_file')} (BM25 Score: {score:.2f})\n{chunk.get('text')}"
            )

        return "\n".join(output)

    def summarize_document(self, file_path: Path, max_length: int = 1200) -> str:
        """Extract a structured summary and highlight outline of a document."""
        raw_text = self.extract_text_from_file(file_path)
        if not raw_text or raw_text.startswith("Error"):
            return raw_text or f"Could not read document '{file_path.name}'."

        words = raw_text.split()
        total_words = len(words)
        preview = " ".join(words[:min(total_words, 250)])
        
        return (
            f"📄 Document Summary for '{file_path.name}':\n"
            f"- File Type: {file_path.suffix.upper()}\n"
            f"- Total Word Count: ~{total_words} words\n"
            f"- Document Preview:\n{preview}..."
        )


rag_engine = DocumentRAGEngine()


@tool(name="search_and_read_documents", description="Extract and answer questions from the content of PDFs, Word docs, CSV, JSON, code, or text files using advanced Okapi BM25 RAG semantic retrieval.")
def search_and_read_documents(query: str, document_name_or_path: Optional[str] = None) -> str:
    """RAG semantic search and content extraction tool."""
    from vision.tools.file_tools import _resolve_user_path

    target_path = None
    if document_name_or_path:
        target_path = _resolve_user_path(document_name_or_path, find_existing_file=True)
    elif working_memory.recent_files:
        for fp in working_memory.recent_files:
            p = Path(fp)
            if p.suffix.lower() in DocumentRAGEngine.SUPPORTED_EXTENSIONS:
                target_path = p
                break

    if not target_path or not target_path.exists():
        return f"Error: Could not locate document '{document_name_or_path or 'recent document'}'. Please specify filename or path."

    working_memory.record_file(str(target_path))
    return rag_engine.search_in_document(target_path, query)


@tool(name="search_documents_in_directory", description="Search across all documents, code, notes, and PDFs in a specified directory using multi-document BM25 semantic RAG.")
def search_documents_in_directory(query: str, directory_path: Optional[str] = None) -> str:
    """Multi-document directory RAG semantic search."""
    from vision.tools.file_tools import _resolve_user_path

    target_dir = None
    if directory_path:
        target_dir = _resolve_user_path(directory_path)
    else:
        # Default to workspace root or user downloads
        target_dir = Path("D:\\VISION")

    if not target_dir or not target_dir.exists():
        return f"Error: Directory '{directory_path or target_dir}' not found."

    return rag_engine.search_in_directory(target_dir, query)


@tool(name="summarize_document", description="Provide a structured summary and preview of a local document, PDF, or code file.")
def summarize_document(document_name_or_path: str) -> str:
    """Generate quick document summary and metadata."""
    from vision.tools.file_tools import _resolve_user_path
    target_path = _resolve_user_path(document_name_or_path, find_existing_file=True)
    if not target_path or not target_path.exists():
        return f"Error: File '{document_name_or_path}' could not be found."
    return rag_engine.summarize_document(target_path)
